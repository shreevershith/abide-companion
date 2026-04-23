"""Convert WRITEUP.md to a formatted Abide-Companion-Writeup.docx."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = Path(__file__).parent.parent / "WRITEUP.md"
DOCX_PATH = Path(__file__).parent.parent / "Abide-Companion-Writeup.docx"

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

AMBER = RGBColor(0xD9, 0x7D, 0x0E)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
CODE_RED = RGBColor(0xC0, 0x39, 0x2B)

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def apply_inline(p, raw, size=11):
    for part in INLINE_RE.split(raw):
        if part.startswith("**") and part.endswith("**"):
            set_font(p.add_run(part[2:-2]), bold=True, size=size)
        elif part.startswith("*") and part.endswith("*"):
            set_font(p.add_run(part[1:-1]), italic=True, size=size)
        elif part.startswith("`") and part.endswith("`"):
            set_font(p.add_run(part[1:-1]), name="Courier New", size=size - 1, color=CODE_RED)
        elif part:
            set_font(p.add_run(part), size=size)


def add_heading(text, level):
    p = doc.add_heading("", level=level)
    p.clear()
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=22, bold=True, color=AMBER)
        p.paragraph_format.space_before = Pt(18)
    elif level == 2:
        set_font(run, size=16, bold=True, color=AMBER)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    elif level == 3:
        set_font(run, size=13, bold=True, color=DARK)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
    else:
        set_font(run, size=11, bold=True, color=DARK)


def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    set_font(p.add_run("\n".join(lines)), name="Courier New", size=9)


def cell_shade(tc, fill_hex):
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def add_table(header_cells, data_rows):
    t = doc.add_table(rows=1, cols=len(header_cells))
    t.style = "Table Grid"
    for i, h in enumerate(header_cells):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].clear()
        set_font(cell.paragraphs[0].add_run(h.strip()), bold=True,
                 color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        cell_shade(cell._tc, "D97D0E")
    for ri, row in enumerate(data_rows):
        fill = "FFFFFF" if ri % 2 == 0 else "FEF3E2"
        cells = t.add_row().cells
        for i, txt in enumerate(row):
            cells[i].paragraphs[0].clear()
            apply_inline(cells[i].paragraphs[0], txt.strip(), size=10)
            cell_shade(cells[i]._tc, fill)
    doc.add_paragraph()


def add_hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D97D0E")
    pBdr.append(bot)
    pPr.append(pBdr)


lines = MD_PATH.read_text(encoding="utf-8").splitlines()
i = 0
in_code = False
code_buf = []

while i < len(lines):
    line = lines[i]

    # fenced code block
    if line.strip().startswith("```"):
        if not in_code:
            in_code  = True
            code_buf = []
        else:
            in_code = False
            add_code_block(code_buf)
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # headings
    if   line.startswith("#### "): add_heading(line[5:].strip(), 4)
    elif line.startswith("### "):  add_heading(line[4:].strip(), 3)
    elif line.startswith("## "):   add_heading(line[3:].strip(), 2)
    elif line.startswith("# "):    add_heading(line[2:].strip(), 1)

    # horizontal rule
    elif line.strip() == "---":
        add_hr()

    # table — collect consecutive | lines
    elif line.startswith("|"):
        tbl_lines = []
        while i < len(lines) and lines[i].startswith("|"):
            tbl_lines.append(lines[i])
            i += 1
        rows = []
        for tl in tbl_lines:
            cells = [c for c in tl.split("|") if c.strip()]
            if any(re.match(r"^[-: ]+$", c.strip()) for c in cells):
                continue
            rows.append(cells)
        if rows:
            add_table(rows[0], rows[1:])
        continue

    # bullet list
    elif re.match(r"^(\s*)[-*] ", line):
        indent = len(line) - len(line.lstrip())
        text   = re.sub(r"^(\s*)[-*] ", "", line)
        p      = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.3 + 0.2 * (indent // 2))
        p.paragraph_format.space_after = Pt(3)
        apply_inline(p, text)

    # numbered list
    elif re.match(r"^\d+\. ", line):
        text = re.sub(r"^\d+\. ", "", line)
        p    = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        apply_inline(p, text)

    # blank line
    elif line.strip() == "":
        pass

    # normal paragraph
    else:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        apply_inline(p, line)

    i += 1

doc.save(str(DOCX_PATH))
print(f"Saved: {DOCX_PATH}")
